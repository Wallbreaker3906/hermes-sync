#!/usr/bin/env python3
"""
Generate 船加网详情图排版任务书 Word document.
Usage: Modify the data sections (products, cases, links) then run.
Requires: pip3 install python-docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# === Page setup: landscape A4 ===
section = doc.sections[0]
section.orientation = 1
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# === Font ===
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# === Helpers ===
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_after = Pt(2)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_after = Pt(12)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

def add_body(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(16)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


# ============================================================
# DATA: Fill in below, then run
# ============================================================

COMPANY = "倍豪船舶"       # 供应商全称
BRAND = "倍豪"             # 品牌简称
PRODUCT_CLASS = "电气系统"  # 一级分类名
TOTAL_SERIES = 3           # 系列数
TOTAL_PAGES = 3            # 详情图张数

# Shared module PDF references (same for all products from this supplier)
SHARED_MODULES = [
    ("企业定位",      "PDF 第 4 页，含公司简介文字 + 插图"),
    ("资质认证",      "PDF 第 6 页，含 CCS/NK/ABS/LR/DNV/BV 认证徽标 + 50+ 专利"),
    ("生产能力",      "PDF 第 25 页，含车间实景 + 生产规模数据"),
    ("应用船型",      "PDF 第 21 页，含民船/军船/特种船舶等应用范围"),
]

# Product series: [品名, 型号, 二级分类, 三级分类, PDF页面]
PRODUCTS = [
    ["液位遥测系统",           "—",    "船用控制设备", "其他设备遥测控制系统",   "p19 上半"],
    ["全船综合控制系统 IVCS",  "IVCS", "船舶自动化",   "船舶综合自动化",         "p19 下半"],
    ["抗横倾系统",             "—",    "船舶自动化",   "移动配载自动控制系统",   "p12"],
]

# Engineering cases: [案例名称, 船东/船厂, 关键供货, 适用SKU]
CASES = [
    ["COSL4 兴旺号\n半潜式深水钻井平台", "烟台中集来福士", "阀门遥控系统 + 液位遥测", "液位遥测系统"],
    ["\"科学号\"\n海洋综合考察船", "中科院海洋研究所\n/ 武昌船舶重工", "液位遥测系统 + 阀门遥控 + 浸水报警", "液位遥测系统"],
    ["71米平台供应船", "中海油 / 浙江华恒\n/ 浙江宏大", "配电板 + 监测报警 + PMS + 阀门遥控 + 紧急停车 + 主机遥控", "全船综合控制 IVCS"],
    ["85米平台供应船", "福建东南造船厂\n/ 广州航通船业", "防横倾系统（3+7船套）", "抗横倾系统"],
]

FALLBACK_TEXT = f"倍豪电气系统产品已服务于中海油、中国科学院、中集来福士等多家客户，涵盖科考船、钻井平台、供应船等多种船型"

# Jump link matrix: [SKU名, 链接1, 链接2, 链接3, 链接4]
JUMP_LINKS = [
    ["液位遥测系统",           "全船综合控制 IVCS", "抗横倾系统", "全回转推进器", "船舶自动化产品"],
    ["全船综合控制 IVCS",      "液位遥测系统",      "抗横倾系统", "全回转推进器", "船舶自动化产品"],
    ["抗横倾系统",             "液位遥测系统",      "全船综合控制 IVCS", "全回转推进器", "船舶自动化产品"],
]

NOTES = [
    f"1. 企业定位、资质认证、生产能力、应用船型 4 个板块与「{COMPANY}动力装置详情图」完全一致，排版时可直接复用。",
    f"2. 电气产品与动力产品为同一供应商（{COMPANY}），品牌和供应简称一致。",
    "3. 排版完成后，跳转链接中涉及的全回转推进器等动力装置产品需链接到对应的动力详情图页面。",
]

# ============================================================
# GENERATION (do not modify below)
# ============================================================

add_title(doc, f'{COMPANY} · {PRODUCT_CLASS}产品详情图排版任务书')
add_subtitle(doc, f'每个 SKU 独立 1 张详情图  |  {TOTAL_SERIES} SKU → {TOTAL_PAGES} 张详情图')

# 一、内容结构
add_heading(doc, '一、详情图内容结构（7 板块）')
add_body(doc, '① 产品图片区 ——每 SKU 独立')
add_body(doc, '② 工程案例 —— 紧跟产品图，每 SKU 独立')
add_body(doc, '③ 企业定位 —— 【共用，p4】所有详情图相同')
add_body(doc, '④ 资质认证 —— 【共用，p6】所有详情图相同')
add_body(doc, '⑤ 生产能力 —— 【共用，p25】所有详情图相同')
add_body(doc, '⑥ 应用船型 —— 【共用，p21】所有详情图相同')
add_body(doc, '⑦ 跳转链接区')

# 二、SKU 对应表
add_heading(doc, f'二、各 SKU 型号与详情页对应表')
add_body(doc, f'{PRODUCT_CLASS}共 {TOTAL_SERIES} 个 SKU，每个独立 1 张详情图，共 {TOTAL_PAGES} 张。')
add_body(doc, f'供应简称：{COMPANY}  |  品牌：{BRAND}')

add_table(doc,
    ['序号', '品名', '型号', '二级分类', '三级分类', '详情图张数', 'PDF 页面'],
    [[str(i+1), *p, '1 张', p[-1]] for i, p in enumerate(PRODUCTS)]
)

# 三、工程案例
add_heading(doc, '三、工程案例匹配')
add_body(doc, f'案例来源：PDF 第 31 页。共选取 {len(CASES)} 个{PRODUCT_CLASS}相关案例。')
add_table(doc,
    ['案例名称', '船东/船厂', '关键供货', '适用 SKU'],
    CASES
)
add_body(doc, f'无匹配案例的 SKU 写：「{FALLBACK_TEXT}」')

# 四、共用模块
add_heading(doc, '四、共用模块缩略图（以下 4 个板块所有详情图完全相同，直接复用）')
for name, desc in SHARED_MODULES:
    add_body(doc, f'板块 {name}', bold=True)
    add_body(doc, desc, indent=True)
    doc.add_paragraph()

# 五、跳转链接
add_heading(doc, '五、跳转链接规则（按 SKU 类型动态调整）')
add_body(doc, '跳转链接穿插位置：链接1→①下方  |  链接2→②下方  |  链接3 & 链接4→⑦')
add_body(doc, '链接完整文字：')
add_body(doc, f'点击查看：{BRAND} Powermaster +XXXX  （XXXX 代指下方表格中的内容）')
add_table(doc,
    ['适用 SKU', '链接 1（最优先）', '链接 2', '链接 3', '链接 4（最后）'],
    JUMP_LINKS
)

# 六、备注
add_heading(doc, '六、备注')
for note in NOTES:
    add_body(doc, note)

# Save
out_path = os.path.expanduser(f'~/Documents/{COMPANY}_{PRODUCT_CLASS}详情图排版任务书.docx')
doc.save(out_path)
print(f'Saved to: {out_path}')
